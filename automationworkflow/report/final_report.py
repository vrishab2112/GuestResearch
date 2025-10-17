from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Dict, List, Tuple, Optional


def _read_json(path: Path) -> Dict:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def _read_jsonl(path: Path) -> List[Dict]:
    if not path.exists():
        return []
    out: List[Dict] = []
    try:
        with path.open("r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    out.append(json.loads(line))
                except Exception:
                    continue
    except Exception:
        return []
    return out


def _call_openai_json(messages: List[Dict], model: str = "gpt-4o") -> Dict:
    import requests
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return {}
    url = "https://api.openai.com/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {
        "model": model,
        "response_format": {"type": "json_object"},
        "messages": messages,
        "temperature": 0.2,
    }
    try:
        r = requests.post(url, headers=headers, json=payload, timeout=180)
        r.raise_for_status()
        content = r.json()["choices"][0]["message"]["content"]
        return json.loads(content)
    except Exception:
        return {}


def _build_about_sections(guest: str, guest_dir: Path, model: str = "gpt-4o") -> Dict:
    """Generate structured About subsections and cache to about_sections.json.

    Schema:
    {
      "timeline": {
        "early_life_background": [str],
        "turning_points": [str],
        "breakthrough_recognition": [str],
        "recent_shifts_focus": [str]
      },
      "life_changing_insights": [{"title": str, "detail": str}]
    }
    """
    out_path = guest_dir / "about_sections.json"
    if out_path.exists():
        try:
            return json.loads(out_path.read_text(encoding="utf-8"))
        except Exception:
            pass

    # Gather compact context from Agent 1 outputs
    about = _read_jsonl(guest_dir / "about_guest.jsonl")
    web = _read_jsonl(guest_dir / "web.jsonl")
    yt = _read_jsonl(guest_dir / "raw" / "youtube.jsonl")

    about_text = (about[0] or {}).get("summary", "") if about else ""
    web_snips: List[Dict] = []
    for r in web[:8]:
        if (r.get("text") or "").strip():
            web_snips.append({"title": r.get("title"), "url": r.get("url"), "text": (r.get("text") or "")[:800]})
    # A few transcript snippets for concrete milestones
    transcripts = [r for r in yt if r.get("source_type") == "youtube_transcript" and r.get("text")]
    for t in transcripts[:4]:
        web_snips.append({"title": "yt_transcript", "url": f"https://www.youtube.com/watch?v={t.get('video_id')}", "text": (t.get("text") or "")[:800]})

    context_json = json.dumps({"about": about_text, "sources": web_snips}, ensure_ascii=False)
    system = (
        "You are a biographical editor. Structure key life milestones and insights from the given context. "
        "Return strict JSON with keys: {\n"
        "  \"timeline\": {\n"
        "    \"early_life_background\": [string],\n"
        "    \"turning_points\": [string],\n"
        "    \"breakthrough_recognition\": [string],\n"
        "    \"recent_shifts_focus\": [string]\n"
        "  },\n"
        "  \"life_changing_insights\": [{title, detail}]\n"
        "}\n"
        "Formatting rules for timeline bullets: Start each item with a year or date when identifiable (YYYY or YYYY-MM). If unknown, use 'n.d.'. Then an en dash and a concise description, e.g., '2018 — Joined X'."
    )
    user = f"Guest: {guest}\n\nContext JSON:\n{context_json}"
    messages = [{"role": "system", "content": system}, {"role": "user", "content": user}]
    result = _call_openai_json(messages, model=model)
    if result:
        try:
            out_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception:
            pass
    return result or {}


def _build_communication_assessment(guest: str, guest_dir: Path, model: str = "gpt-4o") -> Dict:
    """Summarize guest's communication style. Cached to communication.json

    Schema: {style: string, tendencies:[string], watchouts:[string], host_tips:[string]}
    """
    out_path = guest_dir / "communication.json"
    if out_path.exists():
        try:
            return json.loads(out_path.read_text(encoding="utf-8"))
        except Exception:
            pass
    about = _read_jsonl(guest_dir / "about_guest.jsonl")
    plan = _read_json(guest_dir / "agent3" / "plan.json")
    topics = plan.get("topics", []) if plan else []
    questions = plan.get("questions", []) if plan else []
    context = {
        "about": (about[0] or {}).get("summary", "") if about else "",
        "sample_topics": [t.get("title") for t in topics[:6]],
        "sample_questions": [q.get("q") for q in questions[:8]],
    }
    system = (
        "You are a communication analyst. Describe this guest's speaking style for a podcast host. "
        "Return strict JSON: {style: string, tendencies:[string], watchouts:[string], host_tips:[string]}"
    )
    user = f"Guest: {guest}\n\nContext JSON:\n{json.dumps(context, ensure_ascii=False)}"
    messages = [{"role": "system", "content": system}, {"role": "user", "content": user}]
    result = _call_openai_json(messages, model=model)
    if result:
        try:
            out_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception:
            pass
    return result or {}


def _build_podcast_appearances(guest: str, guest_dir: Path, model: str = "gpt-4o") -> List[Dict]:
    """Curate podcast/talk appearances with 1-line context. Cached to appearances.json

    Schema: [{title, url, context}]
    """
    out_path = guest_dir / "appearances.json"
    if out_path.exists():
        try:
            return json.loads(out_path.read_text(encoding="utf-8"))
        except Exception:
            pass
    web = _read_jsonl(guest_dir / "web.jsonl")
    yt = _read_jsonl(guest_dir / "raw" / "youtube.jsonl")
    candidates: List[Dict] = []
    for r in web:
        title = (r.get("title") or "").lower()
        url = r.get("url") or ""
        if any(k in title for k in ("podcast", "interview", "talk")) or any(k in url for k in ("podcast", "interview")):
            candidates.append({"title": r.get("title"), "url": url})
    for v in yt:
        if v.get("source_type") == "youtube_video":
            title = (v.get("title") or "").lower()
            if any(k in title for k in ("podcast", "interview", "experience", "talk")):
                url = v.get("url") or (f"https://www.youtube.com/watch?v={v.get('video_id')}" if v.get("video_id") else "")
                candidates.append({"title": v.get("title"), "url": url})
    context = {"items": candidates[:20]}
    system = (
        "You are a producer. From the provided links, select 5-10 notable podcast/interview/talk appearances. "
        "Return JSON array of {title, url, context} with a one-line why it matters."
    )
    user = json.dumps(context, ensure_ascii=False)
    messages = [{"role": "system", "content": system}, {"role": "user", "content": user}]
    result = _call_openai_json(messages, model=model)
    items = result.get("items") if isinstance(result, dict) else result
    items = items if isinstance(items, list) else []
    try:
        out_path.write_text(json.dumps(items, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass
    return items


def _build_topic_deep_dive(guest: str, guest_dir: Path, model: str = "gpt-4o") -> List[Dict]:
    """Editorial angles distinct from plan topics. Cached to topic_deep_dive.json

    Schema: [{title, blurb}]
    """
    out_path = guest_dir / "topic_deep_dive.json"
    if out_path.exists():
        try:
            return json.loads(out_path.read_text(encoding="utf-8"))
        except Exception:
            pass
    about = _read_jsonl(guest_dir / "about_guest.jsonl")
    plan = _read_json(guest_dir / "agent3" / "plan.json")
    about_text = (about[0] or {}).get("summary", "") if about else ""
    topics = plan.get("topics", []) if plan else []
    context = {"about": about_text, "plan_topics": [t.get("title") for t in topics[:10]]}
    system = (
        "You are an editorial lead. Propose 5-8 sharp 'Topic deep dive' angles distinct from existing plan topics. "
        "Return JSON array of {title, blurb}."
    )
    user = json.dumps(context, ensure_ascii=False)
    messages = [{"role": "system", "content": system}, {"role": "user", "content": user}]
    result = _call_openai_json(messages, model=model)
    items = result.get("angles") if isinstance(result, dict) else result
    items = items if isinstance(items, list) else []
    try:
        out_path.write_text(json.dumps(items, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass
    return items


def _build_shorts_ideas(guest: str, guest_dir: Path, model: str = "gpt-4o") -> List[Dict]:
    """Shorts/Reels prompts and optional links. Cached to shorts_ideas.json

    Schema: [{title, prompt, url?}]
    """
    out_path = guest_dir / "shorts_ideas.json"
    if out_path.exists():
        try:
            return json.loads(out_path.read_text(encoding="utf-8"))
        except Exception:
            pass
    yt = _read_jsonl(guest_dir / "raw" / "youtube.jsonl")
    videos = [r for r in yt if r.get("source_type") == "youtube_video"]
    links = [{"title": v.get("title"), "url": v.get("url") or f"https://www.youtube.com/watch?v={v.get('video_id')}"} for v in videos[:12]]
    context = {"videos": links}
    system = (
        "You are a social producer. Create 6-10 Shorts/Reels ideas with a title and a one-line prompt."
        "If a provided link fits the idea, include it. Return JSON array of {title, prompt, url?}."
    )
    user = json.dumps(context, ensure_ascii=False)
    messages = [{"role": "system", "content": system}, {"role": "user", "content": user}]
    result = _call_openai_json(messages, model=model)
    items = result.get("items") if isinstance(result, dict) else result
    items = items if isinstance(items, list) else []
    try:
        out_path.write_text(json.dumps(items, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass
    return items

def _format_about_sections_markdown(sections: Dict) -> str:
    if not sections:
        return ""
    lines: List[str] = []
    tl = sections.get("timeline") or {}
    if any(tl.values()):
        lines.append("### Timeline (Key Milestones)")
        if tl.get("early_life_background"):
            lines.append("#### Early life & background")
            for b in tl.get("early_life_background", [])[:10]:
                lines.append(f"- {b}")
        if tl.get("turning_points"):
            lines.append("#### Turning points / key decisions")
            for b in tl.get("turning_points", [])[:10]:
                lines.append(f"- {b}")
        if tl.get("breakthrough_recognition"):
            lines.append("#### Breakthrough / recognition")
            for b in tl.get("breakthrough_recognition", [])[:10]:
                lines.append(f"- {b}")
        if tl.get("recent_shifts_focus"):
            lines.append("#### Recent shifts / current focus")
            for b in tl.get("recent_shifts_focus", [])[:10]:
                lines.append(f"- {b}")
        lines.append("")
    lci = sections.get("life_changing_insights") or []
    if lci:
        lines.append("### Life‑changing moments / key insights")
        for it in lci[:6]:
            title = it.get("title") or "Insight"
            detail = it.get("detail") or ""
            lines.append(f"- {title}")
            if detail:
                lines.append(f"  - {detail}")
        lines.append("")
    return "\n".join(lines)


def _build_topics_blocks(guest: str, guest_dir: Path, model: str = "gpt-4o") -> Dict:
    """Produce common topics/stories and unexplored depths from gathered context.

    Schema: {"common_topics": [string], "unexplored_depths": [string]}
    Cached at topics_blocks.json.
    """
    out_path = guest_dir / "topics_blocks.json"
    if out_path.exists():
        try:
            return json.loads(out_path.read_text(encoding="utf-8"))
        except Exception:
            pass
    about = _read_jsonl(guest_dir / "about_guest.jsonl")
    web = _read_jsonl(guest_dir / "web.jsonl")
    yt = _read_jsonl(guest_dir / "raw" / "youtube.jsonl")
    plan = _read_json(guest_dir / "agent3" / "plan.json")
    topics = plan.get("topics", []) if plan else []
    questions = plan.get("questions", []) if plan else []
    comments = [r for r in yt if (r.get("source_type") or "").startswith("youtube_comment")]
    comments = sorted(comments, key=lambda c: int(c.get("like_count", 0) or 0), reverse=True)[:120]

    context = {
        "about": (about[0] or {}).get("summary", "") if about else "",
        "web": [{"title": r.get("title"), "url": r.get("url")} for r in web[:12]],
        "plan_topics": [t.get("title") for t in topics[:10]],
        "plan_questions": [q.get("q") for q in questions[:15]],
        "top_comments": [c.get("text") for c in comments[:60]],
    }
    system = (
        "You are an editorial strategist. From the context, list recurring narratives and fresh angles. "
        "Return strict JSON: {\"common_topics\": [string], \"unexplored_depths\": [string]}"
    )
    user = f"Guest: {guest}\n\nContext JSON:\n{json.dumps(context, ensure_ascii=False)}"
    messages = [{"role": "system", "content": system}, {"role": "user", "content": user}]
    result = _call_openai_json(messages, model=model)
    if result:
        try:
            out_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception:
            pass
    return result or {}

def _summarize_agent1(guest_dir: Path) -> Tuple[str, Dict[str, List[str]], List[str], List[Dict]]:
    about = _read_jsonl(guest_dir / "about_guest.jsonl")
    web = _read_jsonl(guest_dir / "web.jsonl")
    books = _read_jsonl(guest_dir / "books_written.jsonl")
    social = _read_jsonl(guest_dir / "social_bio.jsonl")

    about_text = ""
    if about:
        about_text = (about[0] or {}).get("summary", "")
    about_sources = (about[0] or {}).get("sources", []) if about else []

    # Build grouped link sections for nicer report formatting
    link_sections: Dict[str, List[str]] = {"articles": [], "books": [], "social": [], "youtube": []}
    for r in web[:10]:
        url = r.get("url")
        title = (r.get("title") or url or "").strip()
        if url:
            link_sections["articles"].append(f"- [{title}]({url})")
    # Books/longform
    for r in books[:8]:
        url = r.get("url")
        if url:
            label = "Book" if any(d in url for d in ("amazon.com", "goodreads.com", "books.google", "penguin")) else "Link"
            link_sections["books"].append(f"- [{label}]({url})")
    # Social
    for r in social[:8]:
        url = r.get("url")
        if url:
            link_sections["social"].append(f"- [Social]({url})")

    stats: List[str] = []
    youtube = _read_jsonl(guest_dir / "raw" / "youtube.jsonl")
    # YouTube video links
    yt_videos = [r for r in youtube if r.get("source_type") == "youtube_video"]
    for v in yt_videos[:10]:
        url = v.get("url") or (f"https://www.youtube.com/watch?v={v.get('video_id')}" if v.get("video_id") else None)
        title = (v.get("title") or url or "").strip()
        if url:
            link_sections["youtube"].append(f"- [{title}]({url})")
    videos = len(yt_videos)
    comments = sum(1 for r in youtube if (r.get("source_type") or "").startswith("youtube_comment"))
    transcripts = sum(1 for r in youtube if r.get("source_type") == "youtube_transcript")
    if any([videos, comments, transcripts]):
        stats.append(f"YouTube videos: {videos}")
        stats.append(f"Transcripts: {transcripts}")
        stats.append(f"Top comments: {comments}")

    return about_text, link_sections, stats, about_sources


def _format_north_star(north_star_obj: Dict) -> str:
    lines: List[str] = []
    for i, item in enumerate(north_star_obj.get("north_star", []), start=1):
        title = item.get("title") or f"North Star {i}"
        why = item.get("why_it_matters") or ""
        lines.append(f"#### {i}. {title}")
        if why:
            lines.append(f"{why}")
        evidence = item.get("supporting_evidence") or []
        for ev in evidence[:5]:
            lines.append(f"- {ev}")
        lines.append("")
    lesser = north_star_obj.get("lesser_known", [])
    if lesser:
        lines.append("### Lesser‑known topics")
        for t in lesser:
            lines.append(f"- {t}")
        lines.append("")
    return "\n".join(lines)


def _format_plan(plan_obj: Dict, comment_analysis: Optional[Dict] = None, *, guest: str = "", guest_dir: Optional[Path] = None) -> str:
    lines: List[str] = []
    outline = plan_obj.get("outline", [])
    if outline:
        lines.append("### Conversation outline")
        for section in outline:
            sec = section.get("section") or "Section"
            lines.append(f"- **{sec}**")
            for b in section.get("bullets", [])[:10]:
                lines.append(f"  - {b}")
        lines.append("")

    topics = plan_obj.get("topics", [])
    if topics:
        lines.append("### Discussion topics")
        for i, t in enumerate(topics[:10], start=1):
            title = t.get("title") or f"Topic {i}"
            why = t.get("why") or ""
            lines.append(f"#### {i}. {title}")
            if why:
                lines.append(why)
            for c in t.get("citations", [])[:5]:
                lines.append(f"- {c}")
            lines.append("")

    # Prefer themed questions if provided
    q_by_theme = plan_obj.get("questions_by_theme", [])
    if q_by_theme:
        lines.append("### Questions")
        # Enforce exactly 3 themes with 6–7 questions displayed if available
        for t in q_by_theme[:3]:
            theme = t.get("theme") or "Theme"
            sub = t.get("subtitle") or ""
            lines.append(f"#### {theme}")
            if sub:
                lines.append(f"_{sub}_")
            qs = (t.get("questions") or [])
            # Prefer 6–7
            display_qs = qs[:7] if len(qs) >= 7 else qs[:6] if len(qs) >= 6 else qs[:len(qs)]
            for i, q in enumerate(display_qs, start=1):
                text = q.get("q") or f"Question {i}"
                lines.append(f"- {text}")
                for c in q.get("citations", [])[:3]:
                    lines.append(f"  - {c}")
            lines.append("")
    else:
        # Fallback: group flat questions into 3 themed segments with 6–7 questions each
        questions = plan_obj.get("questions", []) or []
        if questions:
            lines.append("### Questions")
            # Compute chunk sizes aiming for 3 groups of 6–7
            max_total = min(len(questions), 21)
            base_chunk = max_total // 3
            remainder = max_total % 3
            sizes = [base_chunk + (1 if i < remainder else 0) for i in range(3)]
            # Ensure each is at least 6 if possible
            sizes = [max(6, s) for s in sizes]
            # Adjust to not exceed available
            while sum(sizes) > max_total:
                for i in range(3):
                    if sizes[i] > 6 and sum(sizes) > max_total:
                        sizes[i] -= 1
            idx = 0
            for seg_idx, seg_size in enumerate(sizes, start=1):
                seg = questions[idx: idx + seg_size]
                if not seg:
                    break
                lines.append(f"#### Segment {seg_idx}")
                for q in seg:
                    text = q.get("q") or "Question"
                    lines.append(f"- {text}")
                    for c in q.get("citations", [])[:3]:
                        lines.append(f"  - {c}")
                lines.append("")
                idx += seg_size
    # Subsection: Questions the audience still wants to know (from comments)
    if comment_analysis:
        open_q = comment_analysis.get("open_questions") or []
        if open_q:
            lines.append("#### Questions the audience still wants to know")
            for q in open_q[:10]:
                lines.append(f"- {q}")
            lines.append("")

    # Audience psychology
    audience = plan_obj.get("audience_psychology")
    if audience:
        lines.append("### Audience psychology")
        if audience.get("summary"):
            lines.append(audience["summary"]) 
            lines.append("")
        for t in (audience.get("themes") or [])[:8]:
            theme = t.get("theme") or "Theme"
            why = t.get("why") or ""
            lines.append(f"- {theme}: {why}")
            for c in t.get("citations", [])[:3]:
                lines.append(f"  - {c}")
        # Inline comment analysis under audience psychology
        if comment_analysis:
            overall = comment_analysis.get("overall_sentiment") or {}
            topics = comment_analysis.get("hot_topics") or []
            controversies = comment_analysis.get("controversies") or []
            open_q = comment_analysis.get("open_questions") or []
            top = comment_analysis.get("top_comments") or []
            stats = comment_analysis.get("stats") or {}

            lines.append("")
            lines.append("#### Comment analysis")
            if overall:
                summary = overall.get("summary") or ""
                pos = overall.get("positive_pct")
                neu = overall.get("neutral_pct")
                neg = overall.get("negative_pct")
                parts: List[str] = []
                if pos is not None:
                    parts.append(f"positive: {pos}%")
                if neu is not None:
                    parts.append(f"neutral: {neu}%")
                if neg is not None:
                    parts.append(f"negative: {neg}%")
                if summary:
                    if parts:
                        lines.append(f"{summary} (" + ", ".join(parts) + ")")
                    else:
                        lines.append(summary)
            if topics:
                lines.append("- Hot topics:")
                for t in topics[:4]:
                    why = t.get("why") or ""
                    lines.append(f"  - {t.get('topic')}: {why}")
            if controversies:
                lines.append("- Controversies:")
                for c in controversies[:3]:
                    lines.append(f"  - {c.get('issue')}: {c.get('sides')}")
            if open_q:
                lines.append("- Open questions (from viewers):")
                for q in open_q[:5]:
                    lines.append(f"  - {q}")
            if top:
                lines.append("- Top comments:")
                for c in top[:3]:
                    likes = c.get("likes")
                    quote = c.get("quote")
                    lines.append(f"  - {likes} likes – {quote}")
            if stats:
                lines.append(f"_Analyzed {stats.get('sample_size')} of {stats.get('total_comments')} comments._")
            lines.append("")
        else:
            lines.append("")

    # Tensions & Struggles (before Insights)
    ts = plan_obj.get("tensions_struggles") or {}
    top_ch = ts.get("top_challenges") or []
    biggest_fear = ts.get("biggest_fear") or {}
    if top_ch or biggest_fear:
        lines.append("### Tensions & struggles")
        for c in top_ch[:2]:
            lines.append(f"- Challenge: {c.get('challenge')}")
            if c.get("why"):
                lines.append(f"  - Why: {c.get('why')}")
        if biggest_fear.get("fear"):
            lines.append(f"- Biggest fear: {biggest_fear.get('fear')}")
            if biggest_fear.get("why"):
                lines.append(f"  - Why: {biggest_fear.get('why')}")
        lines.append("")

    # Controversy, Vulnerability & Taboo
    cvt = plan_obj.get("controversy_vulnerability_taboo") or {}
    debates = cvt.get("controversial_debates") or []
    shame_q = cvt.get("shameful_questions") or []
    if debates or shame_q:
        lines.append("### Controversy, vulnerability & taboo")
        for d in debates[:2]:
            lines.append(f"- Debate: {d.get('debate')}")
            if d.get("sides"):
                lines.append(f"  - Sides: {d.get('sides')}")
        if shame_q:
            lines.append("- Uncomfortable questions:")
            for q in shame_q[:2]:
                lines.append(f"  - {q}")
        lines.append("")

    # Insights & data focused on North Star
    insights = plan_obj.get("insights_data")
    if insights:
        lines.append("### Insights & data")
        for it in insights[:10]:
            head = it.get("headline") or "Insight"
            det = it.get("detail") or ""
            conf = it.get("confidence") or ""
            badge = f" (confidence: {conf})" if conf else ""
            lines.append(f"- {head}{badge}")
            if det:
                lines.append(f"  - {det}")
            for c in (it.get("citations") or [])[:3]:
                lines.append(f"  - {c}")
        lines.append("")

    # Experiments (end of plan)
    experiments = plan_obj.get("experiments")
    if experiments:
        lines.append("### Experiments")
        for ex in experiments[:4]:
            title = ex.get("title") or "Experiment"
            fmt = ex.get("format") or ""
            why = ex.get("why") or ""
            desc = ex.get("description") or ""
            lines.append(f"- {title}")
            if fmt:
                lines.append(f"  - Format: {fmt}")
            if why:
                lines.append(f"  - Why it works: {why}")
            if desc:
                lines.append(f"  - {desc}")
        lines.append("")

    # Shorts/Reels ideas
    shorts = _build_shorts_ideas(guest, guest_dir) if guest and guest_dir else []
    if shorts:
        lines.append("### Shorts / Reels ideas")
        for it in shorts[:10]:
            title = it.get("title") or "Idea"
            prompt = it.get("prompt") or ""
            url = it.get("url") or ""
            lines.append(f"- {title}")
            if prompt:
                lines.append(f"  - {prompt}")
            if url:
                lines.append(f"  - {url}")
        lines.append("")

    return "\n".join(lines)


def _format_comment_analysis(analysis: Dict) -> str:
    lines: List[str] = []
    if not analysis:
        return ""
    overall = analysis.get("overall_sentiment") or {}
    if overall:
        lines.append("### Overall sentiment")
        summary = overall.get("summary") or ""
        if summary:
            lines.append(summary)
        pos = overall.get("positive_pct")
        neu = overall.get("neutral_pct")
        neg = overall.get("negative_pct")
        parts = []
        if pos is not None:
            parts.append(f"positive: {pos}%")
        if neu is not None:
            parts.append(f"neutral: {neu}%")
        if neg is not None:
            parts.append(f"negative: {neg}%")
        if parts:
            lines.append("(" + ", ".join(parts) + ")")
        lines.append("")

    topics = analysis.get("hot_topics") or []
    if topics:
        lines.append("### Hot topics")
        for t in topics[:8]:
            lines.append(f"- {t.get('topic')}: {t.get('why')}")
            if t.get("example_quote"):
                lines.append(f"  - \"{t.get('example_quote')}\"")
        lines.append("")

    controversies = analysis.get("controversies") or []
    if controversies:
        lines.append("### Controversies")
        for c in controversies[:6]:
            lines.append(f"- {c.get('issue')}: {c.get('sides')}")
            if c.get("example_quote"):
                lines.append(f"  - \"{c.get('example_quote')}\"")
        lines.append("")

    open_q = analysis.get("open_questions") or []
    if open_q:
        lines.append("### Open questions from viewers")
        for q in open_q[:10]:
            lines.append(f"- {q}")
        lines.append("")

    top = analysis.get("top_comments") or []
    if top:
        lines.append("### Top comments (by likes)")
        for c in top[:8]:
            likes = c.get("likes")
            quote = c.get("quote")
            lines.append(f"- {likes} likes – {quote}")
        lines.append("")

    stats = analysis.get("stats") or {}
    if stats:
        lines.append(f"_Analyzed {stats.get('sample_size')} of {stats.get('total_comments')} comments._")
        lines.append("")

    return "\n".join(lines)


def build_markdown_report(guest: str, guest_dir: Path) -> str:
    about_text, link_sections, stats, about_sources = _summarize_agent1(guest_dir)
    # Prefer user-selected North Star if present
    north_star_path = guest_dir / "agent2" / "selected_north_star.json"
    if not north_star_path.exists():
        north_star_path = guest_dir / "agent2" / "north_star.json"
    plan_path = guest_dir / "agent3" / "plan.json"
    north_star_obj = _read_json(north_star_path)
    plan_obj = _read_json(plan_path)
    comment_analysis = _read_json(guest_dir / "agent3" / "comment_analysis.json")
    comment_analysis = _read_json(guest_dir / "agent3" / "comment_analysis.json")

    lines: List[str] = []
    lines.append(f"# Final Interview Research Report — {guest}")
    lines.append("")
    if about_text:
        lines.append("## About the guest")
        lines.append(about_text)
        lines.append("")
        # Optional structured subsections (generate if missing)
        about_sections = _build_about_sections(guest, guest_dir)
        if about_sections:
            lines.append(_format_about_sections_markdown(about_sections))
        if about_sources:
            lines.append("#### Sources for 'About the guest'")
            for s in about_sources[:5]:
                url = s.get("url")
                title = s.get("title") or url
                if url:
                    lines.append(f"- [{title}]({url})")
            lines.append("")
    # Grouped links with dedupe and simple authority ranking
    if any(link_sections.values()):
        def _dedupe(seq: List[str]) -> List[str]:
            seen = set()
            out: List[str] = []
            for s in seq:
                if s in seen:
                    continue
                seen.add(s)
                out.append(s)
            return out
        def _rank(seq: List[str]) -> List[str]:
            authority = (
                "wikipedia.org", "britannica.com", "reuters.com", "bbc.com",
                "nytimes.com", "forbes.com", "cnn.com", "npr.org"
            )
            def score(item: str) -> int:
                url = item.split("(")[-1].rstrip(")")
                for i, dom in enumerate(authority):
                    if dom in url:
                        return -100 + i
                return 0
            return sorted(seq, key=score)
        lines.append("## Key links")
        if link_sections.get("articles"):
            art = _rank(_dedupe(link_sections["articles"]))
            lines.append("### Articles")
            lines.extend(art[:15])
            lines.append("")
        if link_sections.get("books"):
            books = _dedupe(link_sections["books"])[:12]
            lines.append("### Books & longform")
            lines.extend(books)
            lines.append("")
            # Insert Common topics / Unexplored depths after books
            topics_blocks = _build_topics_blocks(guest, guest_dir)
            if topics_blocks:
                if (topics_blocks.get("common_topics") or []):
                    lines.append("### Common topics / stories")
                    for it in (topics_blocks.get("common_topics") or [])[:10]:
                        lines.append(f"- {it}")
                    lines.append("")
                if (topics_blocks.get("unexplored_depths") or []):
                    lines.append("### Unexplored depths")
                    for it in (topics_blocks.get("unexplored_depths") or [])[:10]:
                        lines.append(f"- {it}")
                    lines.append("")
            # Insert Podcast appearances
            appearances = _build_podcast_appearances(guest, guest_dir)
            if appearances:
                lines.append("### Podcast appearances")
                for a in appearances[:10]:
                    title = a.get("title") or "Appearance"
                    url = a.get("url") or ""
                    ctx = a.get("context") or ""
                    if url:
                        lines.append(f"- [{title}]({url})")
                    else:
                        lines.append(f"- {title}")
                    if ctx:
                        lines.append(f"  - {ctx}")
                lines.append("")
        if link_sections.get("social"):
            social = _dedupe(link_sections["social"])[:8]
            lines.append("### Social")
            lines.extend(social)
            lines.append("")
        if link_sections.get("youtube"):
            yt = _dedupe(link_sections["youtube"])[:10]
            lines.append("### YouTube")
            lines.extend(yt)
            lines.append("")
    if stats:
        lines.append("## Data gathering summary")
        for s in stats:
            lines.append(f"- {s}")
        lines.append("")

    # Communication assessment
    comm = _build_communication_assessment(guest, guest_dir)
    if comm:
        lines.append("## Communication assessment")
        if comm.get("style"):
            lines.append(comm.get("style"))
        if comm.get("tendencies"):
            lines.append("### Tendencies")
            for t in (comm.get("tendencies") or [])[:8]:
                lines.append(f"- {t}")
        if comm.get("watchouts"):
            lines.append("### Watchouts")
            for t in (comm.get("watchouts") or [])[:8]:
                lines.append(f"- {t}")
        if comm.get("host_tips"):
            lines.append("### Host tips")
            for t in (comm.get("host_tips") or [])[:8]:
                lines.append(f"- {t}")
        lines.append("")

    if north_star_obj:
        lines.append("## North Star insights")
        lines.append(_format_north_star(north_star_obj))
    else:
        lines.append("## North Star insights")
        lines.append("_Not available. Run Agent 2 to generate insights._\n")

    if plan_obj:
        lines.append("## Conversation plan")
        lines.append(_format_plan(plan_obj, comment_analysis=comment_analysis, guest=guest, guest_dir=guest_dir))
    else:
        lines.append("## Conversation plan")
        lines.append("_Not available. Run Agent 3 to generate the plan._\n")

    return "\n".join(lines)


def generate_final_report(guest: str, outputs_root: Path) -> Path:
    guest_dir = outputs_root / guest
    guest_dir.mkdir(parents=True, exist_ok=True)
    content = build_markdown_report(guest, guest_dir)
    report_path = guest_dir / "final_report.md"
    report_path.write_text(content, encoding="utf-8")
    return report_path


def generate_final_report_docx(guest: str, outputs_root: Path) -> Path:
    # Helper: try Pandoc conversion from Markdown if python-docx is unavailable or save fails
    def _pandoc_fallback() -> Path:
        try:
            import pypandoc  # type: ignore
        except Exception:
            return generate_final_report(guest, outputs_root)
        # Build markdown then convert to docx
        guest_dir_local = outputs_root / guest
        guest_dir_local.mkdir(parents=True, exist_ok=True)
        md_content = build_markdown_report(guest, guest_dir_local)
        md_path = guest_dir_local / "final_report.md"
        md_path.write_text(md_content, encoding="utf-8")
        out_path_local = guest_dir_local / "final_report.docx"
        try:
            try:
                # Try conversion directly
                pypandoc.convert_text(md_content, to="docx", format="md", outputfile=str(out_path_local))
            except OSError:
                # Attempt to auto-download pandoc if missing
                try:
                    pypandoc.download_pandoc()
                except Exception:
                    pass
                pypandoc.convert_text(md_content, to="docx", format="md", outputfile=str(out_path_local))
            return out_path_local if out_path_local.exists() else md_path
        except Exception:
            return md_path

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        # Try Pandoc fallback; if that fails, return Markdown
        return _pandoc_fallback()

    guest_dir = outputs_root / guest
    guest_dir.mkdir(parents=True, exist_ok=True)

    # Build sections separately to preserve headings
    about_text, link_sections, stats, about_sources = _summarize_agent1(guest_dir)
    ns_path = guest_dir / "agent2" / "selected_north_star.json"
    if not ns_path.exists():
        ns_path = guest_dir / "agent2" / "north_star.json"
    north_star_obj = _read_json(ns_path)
    plan_obj = _read_json(guest_dir / "agent3" / "plan.json")

    doc = Document()

    # ---- Document base styles ----
    try:
        # Margins ~0.75in for compactness
        section = doc.sections[0]
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

        # Normal style font/spacing
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        pf = normal.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15
    except Exception:
        pass

    # ---- Helpers ----
    def add_bullet(text: str, indent_level: int = 0):
        p = doc.add_paragraph(text, style="List Bullet")
        if indent_level:
            try:
                p.paragraph_format.left_indent = Inches(0.25 * indent_level)
            except Exception:
                pass
        return p

    def add_numbered(text: str, indent_level: int = 0):
        p = doc.add_paragraph(text, style="List Number")
        if indent_level:
            try:
                p.paragraph_format.left_indent = Inches(0.25 * indent_level)
            except Exception:
                pass
        return p

    def add_hyperlink(paragraph, text: str, url: str):
        try:
            part = paragraph.part
            r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)

            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0563C1')
            rPr.append(color)
            new_run.append(rPr)
            t = OxmlElement('w:t')
            t.text = text
            new_run.append(t)
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
        except Exception:
            paragraph.add_run(f"{text} ({url})")

    def parse_md_link(item: str):
        # Accept forms like "- [Title](https://url)" or "[Title](https://url)"
        s = item.strip()
        if s.startswith("- "):
            s = s[2:].strip()
        if s.startswith("[") and "](" in s and s.endswith(")"):
            try:
                title = s[1:s.index("](")]
                url = s[s.index("](") + 2:].strip("()")
                return title, url
            except Exception:
                return item, None
        return item, None
    doc.add_heading(f"Final Interview Research Report — {guest}", 0)

    if about_text:
        doc.add_heading("About the guest", level=1)
        doc.add_paragraph(about_text)
        # Structured subsections (if generated/available)
        about_sections = _build_about_sections(guest, guest_dir)
        if about_sections:
            tl = (about_sections.get("timeline") or {})
            if any(tl.values()):
                doc.add_heading("Timeline (Key Milestones)", level=2)
                if tl.get("early_life_background"):
                    doc.add_heading("Early life & background", level=3)
                    for b in tl.get("early_life_background", [])[:10]:
                        doc.add_paragraph(f"- {b}")
                if tl.get("turning_points"):
                    doc.add_heading("Turning points / key decisions", level=3)
                    for b in tl.get("turning_points", [])[:10]:
                        doc.add_paragraph(f"- {b}")
                if tl.get("breakthrough_recognition"):
                    doc.add_heading("Breakthrough / recognition", level=3)
                    for b in tl.get("breakthrough_recognition", [])[:10]:
                        doc.add_paragraph(f"- {b}")
                if tl.get("recent_shifts_focus"):
                    doc.add_heading("Recent shifts / current focus", level=3)
                    for b in tl.get("recent_shifts_focus", [])[:10]:
                        doc.add_paragraph(f"- {b}")
            lci = about_sections.get("life_changing_insights") or []
            if lci:
                doc.add_heading("Life‑changing moments / key insights", level=2)
                for it in lci[:6]:
                    title = it.get("title") or "Insight"
                    detail = it.get("detail") or ""
                    doc.add_paragraph(f"- {title}")
                    if detail:
                        doc.add_paragraph(f"  - {detail}")
        if about_sources:
            doc.add_heading("Sources for 'About the guest'", level=2)
            for s in about_sources[:5]:
                url = s.get("url")
                title = s.get("title") or url
                if url:
                    doc.add_paragraph(f"{title}: {url}")

    if any(link_sections.values()):
        doc.add_heading("Key links", level=1)
        if link_sections.get("articles"):
            doc.add_heading("Articles", level=2)
            for l in link_sections["articles"]:
                p = doc.add_paragraph()
                title, url = parse_md_link(l)
                if url:
                    add_hyperlink(p, title, url)
                else:
                    p.add_run(title)
        if link_sections.get("books"):
            doc.add_heading("Books & longform", level=2)
            for l in link_sections["books"]:
                p = doc.add_paragraph()
                title, url = parse_md_link(l)
                if url:
                    add_hyperlink(p, title, url)
                else:
                    p.add_run(title)
            # Insert Common topics / Unexplored depths after books
            topics_blocks = _build_topics_blocks(guest, guest_dir)
            if topics_blocks:
                ct = topics_blocks.get("common_topics") or []
                ud = topics_blocks.get("unexplored_depths") or []
                if ct:
                    doc.add_heading("Common topics / stories", level=2)
                    for it in ct[:10]:
                        add_bullet(it)
                if ud:
                    doc.add_heading("Unexplored depths", level=2)
                    for it in ud[:10]:
                        add_bullet(it)
            # Insert Podcast appearances
            appearances = _build_podcast_appearances(guest, guest_dir)
            if appearances:
                doc.add_heading("Podcast appearances", level=2)
                for a in appearances[:10]:
                    title = a.get("title") or "Appearance"
                    url = a.get("url") or ""
                    ctx = a.get("context") or ""
                    add_bullet(title)
                    if url:
                        p = doc.add_paragraph("    ")
                        add_hyperlink(p, url, url)
                    if ctx:
                        doc.add_paragraph(f"  - {ctx}")
        if link_sections.get("social"):
            doc.add_heading("Social", level=2)
            for l in link_sections["social"]:
                p = doc.add_paragraph()
                title, url = parse_md_link(l)
                if url:
                    add_hyperlink(p, title, url)
                else:
                    p.add_run(title)

    if stats:
        doc.add_heading("Data gathering summary", level=1)
        for s in stats:
            doc.add_paragraph(s)

    # Communication assessment
    comm = _build_communication_assessment(guest, guest_dir)
    if comm:
        doc.add_heading("Communication assessment", level=1)
        if comm.get("style"):
            doc.add_paragraph(comm.get("style"))
        if comm.get("tendencies"):
            doc.add_heading("Tendencies", level=2)
            for t in (comm.get("tendencies") or [])[:8]:
                doc.add_paragraph(f"- {t}")
        if comm.get("watchouts"):
            doc.add_heading("Watchouts", level=2)
            for t in (comm.get("watchouts") or [])[:8]:
                doc.add_paragraph(f"- {t}")
        if comm.get("host_tips"):
            doc.add_heading("Host tips", level=2)
            for t in (comm.get("host_tips") or [])[:8]:
                doc.add_paragraph(f"- {t}")

    doc.add_heading("North Star insights", level=1)
    if north_star_obj:
        for i, item in enumerate(north_star_obj.get("north_star", []), start=1):
            title = item.get("title") or f"North Star {i}"
            why = item.get("why_it_matters") or ""
            doc.add_heading(f"{i}. {title}", level=2)
            if why:
                doc.add_paragraph(why)
            for ev in (item.get("supporting_evidence") or [])[:5]:
                doc.add_paragraph(ev)
    else:
        doc.add_paragraph("Not available. Run Agent 2 to generate insights.")

    doc.add_heading("Conversation plan", level=1)
    if plan_obj:
        outline = plan_obj.get("outline", [])
        if outline:
            doc.add_heading("Conversation outline", level=2)
            for section in outline:
                sec = section.get("section") or "Section"
                doc.add_paragraph(sec)
                for b in section.get("bullets", [])[:10]:
                    doc.add_paragraph(f"- {b}")
        topics = plan_obj.get("topics", [])
        if topics:
            doc.add_heading("Discussion topics", level=2)
            for i, t in enumerate(topics[:10], start=1):
                title = t.get("title") or f"Topic {i}"
                why = t.get("why") or ""
                doc.add_heading(f"{i}. {title}", level=3)
                if why:
                    doc.add_paragraph(why)
                for c in t.get("citations", [])[:5]:
                    p = doc.add_paragraph()
                    title, url = parse_md_link(c)
                    if url:
                        add_hyperlink(p, title, url)
                    else:
                        p.add_run(title)
        # Prefer themed questions if present
        q_by_theme = plan_obj.get("questions_by_theme", [])
        if q_by_theme:
            doc.add_heading("Questions", level=2)
            # Enforce exactly 3 themes with 6–7 questions displayed if available
            for t in q_by_theme[:3]:
                theme = t.get("theme") or "Theme"
                sub = t.get("subtitle") or ""
                doc.add_heading(theme, level=3)
                if sub:
                    doc.add_paragraph(sub)
                qs = (t.get("questions") or [])
                display_qs = qs[:7] if len(qs) >= 7 else qs[:6] if len(qs) >= 6 else qs[:len(qs)]
                for idx, q in enumerate(display_qs, start=1):
                    text = q.get("q") or "Question"
                    add_numbered(text)
                    for c in q.get("citations", [])[:3]:
                        p = doc.add_paragraph("    ")
                        title, url = parse_md_link(c)
                        if url:
                            add_hyperlink(p, title, url)
                        else:
                            p.add_run(title)
        else:
            questions = plan_obj.get("questions", [])
            if questions:
                doc.add_heading("Questions", level=2)
                for i, q in enumerate(questions[:25], start=1):
                    text = q.get("q") or f"Question {i}"
                    doc.add_paragraph(f"{i}. {text}")
                    for c in q.get("citations", [])[:3]:
                        doc.add_paragraph(f"  - {c}")
            # Subsection from comment analysis
            comment_analysis = _read_json(guest_dir / "agent3" / "comment_analysis.json")
            if comment_analysis:
                open_q = comment_analysis.get("open_questions") or []
                if open_q:
                    doc.add_heading("Questions the audience still wants to know", level=3)
                    for q in open_q[:10]:
                        doc.add_paragraph(f"- {q}")
        # Audience psychology
        audience = plan_obj.get("audience_psychology")
        if audience:
            doc.add_heading("Audience psychology", level=2)
            if audience.get("summary"):
                doc.add_paragraph(audience["summary"]) 
            for t in (audience.get("themes") or [])[:8]:
                theme = t.get("theme") or "Theme"
                why = t.get("why") or ""
                doc.add_paragraph(f"- {theme}: {why}")
                for c in t.get("citations", [])[:3]:
                    doc.add_paragraph(f"  - {c}")
        # Insights & data
        insights = plan_obj.get("insights_data")
        if insights:
            doc.add_heading("Insights & data", level=2)
            for it in insights[:10]:
                head = it.get("headline") or "Insight"
                det = it.get("detail") or ""
                conf = it.get("confidence") or ""
                badge = f" (confidence: {conf})" if conf else ""
                add_bullet(f"{head}{badge}")
                if det:
                    doc.add_paragraph(f"  - {det}")
                for c in (it.get("citations") or [])[:3]:
                    p = doc.add_paragraph("    ")
                    title, url = parse_md_link(c)
                    if url:
                        add_hyperlink(p, title, url)
                    else:
                        p.add_run(title)
        # Experiments
        experiments = plan_obj.get("experiments")
        if experiments:
            doc.add_heading("Experiments", level=2)
            for ex in experiments[:4]:
                title = ex.get("title") or "Experiment"
                fmt = ex.get("format") or ""
                why = ex.get("why") or ""
                desc = ex.get("description") or ""
                add_bullet(title)
                if fmt:
                    doc.add_paragraph(f"  - Format: {fmt}")
                if why:
                    doc.add_paragraph(f"  - Why it works: {why}")
                if desc:
                    doc.add_paragraph(f"  - {desc}")
        # Shorts/Reels ideas
        shorts = _build_shorts_ideas(guest, guest_dir)
        if shorts:
            doc.add_heading("Shorts / Reels ideas", level=2)
            for it in shorts[:10]:
                title = it.get("title") or "Idea"
                prompt = it.get("prompt") or ""
                url = it.get("url") or ""
                add_bullet(title)
                if prompt:
                    doc.add_paragraph(f"  - {prompt}")
                if url:
                    p = doc.add_paragraph("    ")
                    add_hyperlink(p, url, url)
    else:
        doc.add_paragraph("Not available. Run Agent 3 to generate the plan.")

    # Comment analysis (if present)
    comment_analysis = _read_json(guest_dir / "agent3" / "comment_analysis.json")
    if comment_analysis:
        doc.add_heading("Comment analysis", level=1)
        overall = comment_analysis.get("overall_sentiment") or {}
        if overall:
            doc.add_heading("Overall sentiment", level=2)
            if overall.get("summary"):
                doc.add_paragraph(overall["summary"]) 
            parts = []
            if overall.get("positive_pct") is not None:
                parts.append(f"positive: {overall.get('positive_pct')}%")
            if overall.get("neutral_pct") is not None:
                parts.append(f"neutral: {overall.get('neutral_pct')}%")
            if overall.get("negative_pct") is not None:
                parts.append(f"negative: {overall.get('negative_pct')}%")
            if parts:
                doc.add_paragraph("(" + ", ".join(parts) + ")")
        topics = comment_analysis.get("hot_topics") or []
        if topics:
            doc.add_heading("Hot topics", level=2)
            for t in topics[:8]:
                doc.add_paragraph(f"- {t.get('topic')}: {t.get('why')}")
                if t.get("example_quote"):
                    doc.add_paragraph(f"  - \"{t.get('example_quote')}\"")
        controversies = comment_analysis.get("controversies") or []
        if controversies:
            doc.add_heading("Controversies", level=2)
            for c in controversies[:6]:
                doc.add_paragraph(f"- {c.get('issue')}: {c.get('sides')}")
                if c.get("example_quote"):
                    doc.add_paragraph(f"  - \"{c.get('example_quote')}\"")
        open_q = comment_analysis.get("open_questions") or []
        if open_q:
            doc.add_heading("Open questions from viewers", level=2)
            for q in open_q[:10]:
                doc.add_paragraph(f"- {q}")
        top = comment_analysis.get("top_comments") or []
        if top:
            doc.add_heading("Top comments (by likes)", level=2)
            for c in top[:8]:
                doc.add_paragraph(f"- {c.get('likes')} likes – {c.get('quote')}")
        stats = comment_analysis.get("stats") or {}
        if stats:
            doc.add_paragraph(f"Analyzed {stats.get('sample_size')} of {stats.get('total_comments')} comments.")

    out_path = guest_dir / "final_report.docx"
    # Ensure correct .docx write with fallback
    try:
        doc.save(str(out_path))
        return out_path
    except Exception:
        # Try Pandoc as a secondary fallback; then Markdown as last resort
        alt = _pandoc_fallback()
        return alt


